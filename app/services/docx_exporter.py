import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from slugify import slugify

class DocxExporter:
    """Clase para manejar la exportación de libros a formato DOCX optimizado para Kindle"""
    
    def __init__(self, book):
        """
        Inicializa el exportador con un libro específico
        
        Args:
            book: Instancia del modelo Book a exportar
        """
        self.book = book
        self.document = Document()
        self._setup_document_properties()
    
    def _setup_document_properties(self):
        """Configurar las propiedades del documento para Kindle"""
        # Configurar metadatos
        core_properties = self.document.core_properties
        core_properties.title = self.book.title
        core_properties.subject = self.book.market_niche
        core_properties.category = self.book.market_niche
        
        # Limitar el campo purpose a 255 caracteres para evitar errores
        if self.book.purpose and len(self.book.purpose) > 250:
            purpose_truncated = self.book.purpose[:250] + "..."
            core_properties.comments = purpose_truncated
        else:
            core_properties.comments = self.book.purpose
        
        # Configurar estilos
        self._setup_styles()
        
        # Configurar tamaño de página (optimizado para Kindle)
        section = self.document.sections[0]
        section.page_width = Inches(6)
        section.page_height = Inches(9)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    def _setup_styles(self):
        """Configurar los estilos del documento para Kindle"""
        styles = self.document.styles
        
        # Estilo de título principal
        title_style = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
        title_format = title_style.font
        title_format.name = 'Calibri'
        title_format.size = Pt(28)
        title_format.bold = True
        title_paragraph_format = title_style.paragraph_format
        title_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph_format.space_before = Pt(0)
        title_paragraph_format.space_after = Pt(20)
        
        # Estilo de capítulo
        chapter_style = styles.add_style('ChapterStyle', WD_STYLE_TYPE.PARAGRAPH)
        chapter_format = chapter_style.font
        chapter_format.name = 'Calibri'
        chapter_format.size = Pt(24)
        chapter_format.bold = True
        chapter_paragraph_format = chapter_style.paragraph_format
        chapter_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chapter_paragraph_format.space_before = Pt(24)
        chapter_paragraph_format.space_after = Pt(12)
        chapter_paragraph_format.page_break_before = True  # Cada capítulo comienza en nueva página
        
        # Estilo de subtítulo
        heading_style = styles.add_style('HeadingStyle', WD_STYLE_TYPE.PARAGRAPH)
        heading_format = heading_style.font
        heading_format.name = 'Calibri'
        heading_format.size = Pt(16)
        heading_format.bold = True
        heading_paragraph_format = heading_style.paragraph_format
        heading_paragraph_format.space_before = Pt(16)
        heading_paragraph_format.space_after = Pt(8)
        
        # Estilo de párrafo normal
        normal_style = styles['Normal']
        normal_format = normal_style.font
        normal_format.name = 'Calibri'
        normal_format.size = Pt(12)
        normal_paragraph_format = normal_style.paragraph_format
        normal_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        normal_paragraph_format.space_after = Pt(10)
        normal_paragraph_format.first_line_indent = Pt(20)
        
        # Crear estilos para la tabla de contenidos
        # TOC 1 - para entradas principales (capítulos)
        toc1_style = styles.add_style('TOC 1', WD_STYLE_TYPE.PARAGRAPH)
        toc1_format = toc1_style.font
        toc1_format.name = 'Calibri'
        toc1_format.size = Pt(12)
        toc1_format.bold = True
        toc1_paragraph_format = toc1_style.paragraph_format
        toc1_paragraph_format.space_after = Pt(6)
        toc1_paragraph_format.left_indent = Inches(0)
        
        # TOC 2 - para entradas secundarias (subtítulos)
        toc2_style = styles.add_style('TOC 2', WD_STYLE_TYPE.PARAGRAPH)
        toc2_format = toc2_style.font
        toc2_format.name = 'Calibri'
        toc2_format.size = Pt(11)
        toc2_format.italic = True
        toc2_paragraph_format = toc2_style.paragraph_format
        toc2_paragraph_format.space_after = Pt(6)
        toc2_paragraph_format.left_indent = Inches(0.25)
    
    def _add_bookmark(self, paragraph, bookmark_name):
        """Añade un marcador (bookmark) a un párrafo, útil para la navegación en Kindle"""
        run = paragraph.add_run()
        tag = run._element
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), '0')
        start.set(qn('w:name'), bookmark_name)
        tag.append(start)
        
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), '0')
        tag.append(end)
    
    def _extract_headings(self, text):
        """
        Extrae los encabezados del texto del capítulo.
        Busca líneas que parezcan encabezados (generalmente en negritas o con formato especial).
        """
        # Patrón para detectar encabezados (líneas cortas que no terminan con punto)
        heading_pattern = r"(?m)^([A-Z][^\.]{5,60})$"
        
        # Detectar posibles encabezados
        potential_headings = re.findall(heading_pattern, text)
        
        # Filtrar para evitar falsos positivos
        headings = []
        for heading in potential_headings:
            # Verificar que no sea parte de un párrafo (tiene líneas vacías alrededor)
            # y que sea lo suficientemente corto para ser un encabezado
            if heading.strip() and len(heading.strip()) < 80:
                headings.append(heading.strip())
        
        return headings
    
    def _process_chapter_content(self, content):
        """
        Procesa el contenido del capítulo para separar encabezados y párrafos.
        
        Args:
            content: Texto del capítulo
            
        Returns:
            Lista de tuplas (tipo, texto) donde tipo puede ser 'heading' o 'paragraph'
        """
        processed_content = []
        
        # Dividir el contenido por líneas
        lines = content.strip().split("\n")
        
        i = 0
        while i < len(lines):
            current_line = lines[i].strip()
            
            # Saltar líneas vacías
            if not current_line:
                i += 1
                continue
            
            # Detectar si es un posible encabezado
            is_heading = (
                current_line and 
                len(current_line) < 80 and 
                not current_line.endswith('.') and
                (i == 0 or not lines[i-1].strip()) and
                (i == len(lines)-1 or not lines[i+1].strip() or len(lines[i+1].strip()) > 150)
            )
            
            if is_heading:
                processed_content.append(('heading', current_line))
                i += 1
            else:
                # Recolectar párrafo (líneas consecutivas que no están vacías)
                paragraph_lines = [current_line]
                j = i + 1
                while j < len(lines) and lines[j].strip():
                    paragraph_lines.append(lines[j].strip())
                    j += 1
                
                paragraph_text = ' '.join(paragraph_lines)
                processed_content.append(('paragraph', paragraph_text))
                i = j
        
        return processed_content
    
    def _add_toc_entry(self, text, level=1):
        """Añade una entrada a la tabla de contenidos"""
        paragraph = self.document.add_paragraph()
        if level == 1:
            paragraph.style = 'TOC 1'
        else:
            paragraph.style = 'TOC 2'
            # No es necesario añadir sangría aquí porque ya está definida en el estilo
        
        paragraph.add_run(text)
    
    def generate_docx(self):
        """
        Genera el documento DOCX para el libro.
        
        Returns:
            BytesIO: Contenido del documento DOCX como objeto BytesIO
        """
        try:
            # Agregar página de título
            self._add_title_page()
            
            # Agregar tabla de contenidos
            self._add_table_of_contents()
            
            # Agregar capítulos
            for chapter in sorted(self.book.chapters, key=lambda x: x.chapter_number):
                self._add_chapter(chapter)
            
            # Guardar el documento en un buffer BytesIO
            docx_buffer = io.BytesIO()
            self.document.save(docx_buffer)
            docx_buffer.seek(0)
            
            return docx_buffer
        except Exception as e:
            # Capturar cualquier excepción para facilitar la depuración
            import logging
            logging.error(f"Error al generar DOCX: {str(e)}")
            raise
    
    def _add_title_page(self):
        """Añade la página de título al documento"""
        # Título del libro
        title_paragraph = self.document.add_paragraph(self.book.title, style='TitleStyle')
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salto de página
        self.document.add_page_break()
    
    def _add_table_of_contents(self):
        """Añade la tabla de contenidos al documento"""
        toc_title = self.document.add_paragraph("Tabla de Contenidos", style='TitleStyle')
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar entradas de la tabla de contenidos
        for chapter in sorted(self.book.chapters, key=lambda x: x.chapter_number):
            bookmark_name = f"chapter_{chapter.chapter_number}"
            self._add_toc_entry(f"Capítulo {chapter.chapter_number}: {chapter.title}")
            
            # Extraer subtítulos del capítulo
            headings = self._extract_headings(chapter.content)
            for heading in headings[:5]:  # Limitar a 5 subtítulos por capítulo para no sobrecargar el TOC
                subheading_bookmark = slugify(heading)
                self._add_toc_entry(heading, level=2)
        
        # Salto de página después de la tabla de contenidos
        self.document.add_page_break()
    
    def _add_chapter(self, chapter):
        """
        Añade un capítulo al documento.
        
        Args:
            chapter: Instancia del modelo Chapter
        """
        # Título del capítulo
        chapter_title = f"Capítulo {chapter.chapter_number}: {chapter.title}"
        chapter_para = self.document.add_paragraph(chapter_title, style='ChapterStyle')
        self._add_bookmark(chapter_para, f"chapter_{chapter.chapter_number}")
        
        # Procesar el contenido del capítulo
        processed_content = self._process_chapter_content(chapter.content)
        
        # Añadir el contenido procesado
        for content_type, text in processed_content:
            if content_type == 'heading':
                heading_para = self.document.add_paragraph(text, style='HeadingStyle')
                self._add_bookmark(heading_para, slugify(text))
            else:
                # Es un párrafo normal
                self.document.add_paragraph(text)